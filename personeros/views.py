from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Count, Q
from django.utils import timezone
from django.http import JsonResponse, HttpResponse
from django.conf import settings

from .models import Personero, PerfilUsuario, CentroVotacion, Departamento, Provincia, Distrito, ActaElectoral
from .forms import PersoneroSelfUpdateForm, PersoneroPublicRegistrationForm, ActaElectoralForm


# ── Auth & Registro Público ───────────────────────────────────────────────────

def registro_publico_view(request):
    """Vista pública para que los personeros se registren ellos mismos."""
    if request.method == 'POST':
        form = PersoneroPublicRegistrationForm(request.POST)
        if form.is_valid():
            personero = form.save(commit=False)
            personero.estado = 'pendiente'  # Por defecto queda pendiente de asignación por admin
            personero.save()
            
            # Si el usuario que está registrando ya está autenticado (admin/coordinador),
            # lo redirigimos de vuelta al listado de personeros con un mensaje de éxito.
            if request.user.is_authenticated:
                messages.success(request, f'¡Personero {personero.nombre_completo} registrado exitosamente!')
                return redirect('personeros:lista')
                
            messages.success(request, '¡Gracias por registrarte! Un coordinador se pondrá en contacto contigo pronto para asignarte un local.')
            return render(request, 'personeros/registro_exitoso.html', {'personero': personero})
    else:
        form = PersoneroPublicRegistrationForm()

    # Si es administrador, cargamos el formulario dentro del diseño del panel con menú lateral
    template_name = 'personeros/registro_admin.html' if request.user.is_authenticated else 'personeros/registro_publico.html'
    return render(request, template_name, {'form': form})


# ── Helpers ───────────────────────────────────────────────────────────────────

def get_perfil(user):
    """Retorna el PerfilUsuario o None si no tiene."""
    try:
        return user.perfil
    except PerfilUsuario.DoesNotExist:
        return None


def get_personero_queryset(user):
    """Filtra personeros según el rol del usuario."""
    perfil = get_perfil(user)
    if not perfil or perfil.es_superadmin:
        return Personero.objects.all()
    if perfil.rol == 'coordinador_departamental' and perfil.departamento:
        return Personero.objects.filter(distrito__provincia__departamento=perfil.departamento)
    if perfil.rol == 'coordinador_provincial' and perfil.provincia:
        return Personero.objects.filter(distrito__provincia=perfil.provincia)
    return Personero.objects.none()


def _get_personero_destination(user):
    """Retorna la URL correcta para el personero según el avance de su registro."""
    try:
        personero = user.personero
        if personero.perfil_completado and personero.estado == 'confirmado':
            return 'personeros:reportar_escrutinio'
    except Personero.DoesNotExist:
        pass
    return 'personeros:mi_perfil'


# ── Auth ──────────────────────────────────────────────────────────────────────

def login_view(request):
    if request.user.is_authenticated:
        return _redirect_after_login(request.user)

    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '').strip()
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            return _redirect_after_login(user)
        messages.error(request, 'DNI o contraseña incorrectos.')

    return render(request, 'personeros/login.html')


def _redirect_after_login(user):
    perfil = get_perfil(user)
    if perfil and perfil.es_personero:
        return redirect(_get_personero_destination(user))
    return redirect('personeros:dashboard')


def logout_view(request):
    logout(request)
    return redirect('bienvenida')


# ── Dashboard (coordinadores / admin) ─────────────────────────────────────────

@login_required(login_url='personeros:login')
def dashboard_view(request):
    perfil = get_perfil(request.user)
    if perfil and perfil.es_personero:
        return redirect(_get_personero_destination(request.user))

    qs = get_personero_queryset(request.user)

    # KPIs
    total       = qs.count()
    confirmados = qs.filter(estado='confirmado').count()
    pendientes  = qs.filter(estado='pendiente').count()
    retirados   = qs.filter(estado='retirado').count()
    completados = qs.filter(perfil_completado=True).count()

    pct_confirmados = round((confirmados / total * 100) if total else 0)
    pct_completados = round((completados / total * 100) if total else 0)

    # --- Consolidación de Conteo de Votos (Escrutinio Rápido) ---
    from django.db.models import Sum
    actas_qs = ActaElectoral.objects.filter(personero__in=qs)
    
    votos_totales = actas_qs.aggregate(
        jp=Sum('votos_jp'),
        k=Sum('votos_k'),
        blanco=Sum('votos_blanco'),
        nulos=Sum('votos_nulos'),
        viciados=Sum('votos_viciados'),
    )

    votos_jp = votos_totales['jp'] or 0
    votos_k = votos_totales['k'] or 0
    votos_blanco = votos_totales['blanco'] or 0
    votos_nulos = votos_totales['nulos'] or 0
    votos_viciados = votos_totales['viciados'] or 0

    votos_validos = votos_jp + votos_k
    votos_emitidos = votos_validos + votos_blanco + votos_nulos + votos_viciados

    pct_jp = round((votos_jp / votos_validos * 100) if votos_validos else 0, 1)
    pct_k = round((votos_k / votos_validos * 100) if votos_validos else 0, 1)

    pct_blanco = round((votos_blanco / votos_emitidos * 100) if votos_emitidos else 0, 1)
    pct_nulos = round((votos_nulos / votos_emitidos * 100) if votos_emitidos else 0, 1)
    pct_viciados = round((votos_viciados / votos_emitidos * 100) if votos_emitidos else 0, 1)

    # Avance de Mesas
    mesas_esperadas = qs.filter(estado='confirmado').exclude(numero_mesa='').exclude(centro_votacion__isnull=True).count()
    mesas_procesadas = actas_qs.count()
    pct_mesas = round((mesas_procesadas / mesas_esperadas * 100) if mesas_esperadas else 0, 1)

    actas_recientes = actas_qs.order_by('-fecha_registro')[:10]

    # Por departamento
    por_dpto = (
        qs.values('distrito__provincia__departamento__nombre', 'distrito__provincia__departamento__id_ubigeo')
          .annotate(total=Count('id'), confirmados=Count('id', filter=Q(estado='confirmado')))
          .order_by('-total')
    )
    por_dpto_list = [
        {
            'codigo':      d['distrito__provincia__departamento__id_ubigeo'],
            'nombre':      d['distrito__provincia__departamento__nombre'] or 'Sin asignar',
            'total':       d['total'],
            'confirmados': d['confirmados'],
            'pct':         round(d['confirmados'] / d['total'] * 100) if d['total'] else 0,
        }
        for d in por_dpto
    ]

    # Últimos registros
    recientes = qs.order_by('-fecha_creacion')[:10]

    context = {
        'perfil':          perfil,
        'total':           total,
        'confirmados':     confirmados,
        'pendientes':      pendientes,
        'retirados':       retirados,
        'completados':     completados,
        'pct_confirmados': pct_confirmados,
        'pct_completados': pct_completados,
        'por_dpto':        por_dpto_list,
        'recientes':       recientes,
        
        # Conteo
        'votos_jp':        votos_jp,
        'votos_k':         votos_k,
        'votos_blanco':    votos_blanco,
        'votos_nulos':     votos_nulos,
        'votos_viciados':  votos_viciados,
        'votos_validos':   votos_validos,
        'votos_emitidos':  votos_emitidos,
        'pct_jp':          pct_jp,
        'pct_k':           pct_k,
        'pct_blanco':      pct_blanco,
        'pct_nulos':       pct_nulos,
        'pct_viciados':    pct_viciados,
        'mesas_esperadas': mesas_esperadas,
        'mesas_procesadas':mesas_procesadas,
        'pct_mesas':       pct_mesas,
        'actas_recientes': actas_recientes,
    }
    return render(request, 'personeros/dashboard.html', context)


# ── Lista de personeros ────────────────────────────────────────────────────────

@login_required(login_url='personeros:login')
def lista_view(request):
    perfil = get_perfil(request.user)
    if perfil and perfil.es_personero:
        return redirect(_get_personero_destination(request.user))

    qs = get_personero_queryset(request.user)

    # Filtros
    q          = request.GET.get('q', '')
    estado     = request.GET.get('estado', '')
    dpto_id    = request.GET.get('departamento', '')

    if q:
        qs = qs.filter(
            Q(nombres__icontains=q) |
            Q(apellido_paterno__icontains=q) |
            Q(apellido_materno__icontains=q) |
            Q(dni__icontains=q) |
            Q(nro_celular__icontains=q)
        )
    if estado:
        qs = qs.filter(estado=estado)
    if dpto_id:
        qs = qs.filter(distrito__provincia__departamento__id_ubigeo=dpto_id)

    context = {
        'perfil':       perfil,
        'personeros':   qs.order_by('apellido_paterno'),
        'departamentos': Departamento.objects.all(),
        'q':            q,
        'estado_sel':   estado,
        'dpto_sel':     dpto_id,
        'total':        qs.count(),
    }
    return render(request, 'personeros/lista.html', context)


# ── Mi Perfil (personero auto-actualización) ──────────────────────────────────

@login_required(login_url='personeros:login')
def mi_perfil_view(request):
    try:
        personero = request.user.personero
    except Personero.DoesNotExist:
        from django.contrib.auth import logout
        logout(request)
        messages.error(request, 'No tienes un perfil de personero asociado.')
        return redirect('personeros:login')

    if request.method == 'POST':
        # Guardar perfil
        if personero.perfil_completado:
            messages.warning(request, 'Tu perfil ya fue completado y no puede modificarse.')
            return redirect('personeros:mi_perfil')

        form = PersoneroSelfUpdateForm(request.POST, instance=personero)
        if form.is_valid():
            p = form.save(commit=False)
            p.perfil_completado = True
            p.fecha_completado  = timezone.now()
            p.estado            = 'confirmado'
            p.save()
            messages.success(request, '¡Tus datos fueron guardados exitosamente!')
            return redirect('personeros:mi_perfil')
    else:
        form = PersoneroSelfUpdateForm(instance=personero)

    centro_preasignado = None
    if not personero.perfil_completado and personero.centro_votacion:
        centro_preasignado = personero.centro_votacion

    context = {
        'personero':         personero,
        'form':              form,
        'readonly':          personero.perfil_completado,
        'centro_preasignado': centro_preasignado,
    }
    return render(request, 'personeros/mi_perfil.html', context)


# ── Reportar Escrutinio por Mesa ──────────────────────────────────────────────

@login_required(login_url='personeros:login')
def reportar_escrutinio_view(request):
    try:
        personero = request.user.personero
    except Personero.DoesNotExist:
        from django.contrib.auth import logout
        logout(request)
        messages.error(request, 'No tienes un perfil de personero asociado.')
        return redirect('personeros:login')

    actas_reportadas = personero.actas.all().order_by('-fecha_registro')
    form_acta = None

    if request.method == 'POST':
        if not personero.perfil_completado:
            messages.error(request, 'Primero debes completar tus datos electorales.')
            return redirect('personeros:reportar_escrutinio')
        if not personero.centro_votacion or not personero.numero_mesa:
            messages.error(request, 'Debes tener un local y una mesa asignados para reportar.')
            return redirect('personeros:reportar_escrutinio')

        form_acta = ActaElectoralForm(request.POST, request.FILES)
        if form_acta.is_valid():
            numero_mesa = form_acta.cleaned_data.get('numero_mesa', '').strip()
            
            # Validar si esta mesa ya fue reportada en este local de votación
            if ActaElectoral.objects.filter(centro_votacion=personero.centro_votacion, numero_mesa=numero_mesa).exists():
                messages.error(request, f'La mesa {numero_mesa} de tu local ya fue reportada.')
                # Volver a cargar el formulario con el error
                context = {
                    'personero':        personero,
                    'actas_reportadas': actas_reportadas,
                    'form_acta':        form_acta,
                }
                return render(request, 'personeros/reportar_escrutinio.html', context)

            a = form_acta.save(commit=False)
            a.personero = personero
            a.centro_votacion = personero.centro_votacion
            a.save()
            messages.success(request, f'¡Resultados de la mesa {numero_mesa} guardados con éxito!')
            return redirect('personeros:reportar_escrutinio')
        else:
            messages.error(request, 'Hubo un error al validar los datos del conteo de votos.')
    else:
        # Pre-rellenar con la mesa preasignada al personero si no ha reportado nada aún
        initial_data = {}
        if personero.numero_mesa and not actas_reportadas.filter(numero_mesa=personero.numero_mesa).exists():
            initial_data['numero_mesa'] = personero.numero_mesa
        form_acta = ActaElectoralForm(initial=initial_data)

    context = {
        'personero':        personero,
        'actas_reportadas': actas_reportadas,
        'form_acta':        form_acta,
    }
    return render(request, 'personeros/reportar_escrutinio.html', context)



# ── APIs para selectores en cascada ───────────────────────────────────────────

def api_provincias_view(request):
    depto_id = request.GET.get('departamento')
    if not depto_id:
        return JsonResponse({'provincias': []})
    provincias = Provincia.objects.filter(departamento_id=depto_id).values('id_ubigeo', 'nombre')
    return JsonResponse({'provincias': list(provincias)})


def api_distritos_view(request):
    prov_id = request.GET.get('provincia')
    if not prov_id:
        return JsonResponse({'distritos': []})
    distritos = Distrito.objects.filter(provincia_id=prov_id).values('id_ubigeo', 'nombre')
    return JsonResponse({'distritos': list(distritos)})


def api_centros_view(request):
    dist_id = request.GET.get('distrito')
    if not dist_id:
        return JsonResponse({'centros': []})
    centros = CentroVotacion.objects.filter(distrito_id=dist_id).values('id', 'nombre', 'direccion')
    return JsonResponse({'centros': list(centros)})


def api_resumen_view(request):
    qs = get_personero_queryset(request.user)
    por_dpto = (
        qs.values('distrito__provincia__departamento__nombre', 'distrito__provincia__departamento__id_ubigeo')
          .annotate(total=Count('id'), confirmados=Count('id', filter=Q(estado='confirmado')))
    )
    data = [
        {
            'codigo':      d['distrito__provincia__departamento__id_ubigeo'],
            'nombre':      d['distrito__provincia__departamento__nombre'] or 'Sin asignar',
            'total':       d['total'],
            'confirmados': d['confirmados'],
            'pct':         round(d['confirmados'] / d['total'] * 100) if d['total'] else 0,
        }
        for d in por_dpto if d['distrito__provincia__departamento__id_ubigeo']
    ]
    return JsonResponse({'departamentos': data})


@login_required(login_url='personeros:login')
def exportar_excel_view(request):
    perfil = get_perfil(request.user)
    if perfil and perfil.es_personero:
        return redirect(_get_personero_destination(request.user))

    qs = get_personero_queryset(request.user)

    # Filtros
    q          = request.GET.get('q', '')
    estado     = request.GET.get('estado', '')
    dpto_id    = request.GET.get('departamento', '')

    if q:
        qs = qs.filter(
            Q(nombres__icontains=q) |
            Q(apellido_paterno__icontains=q) |
            Q(apellido_materno__icontains=q) |
            Q(dni__icontains=q) |
            Q(nro_celular__icontains=q)
        )
    if estado:
        qs = qs.filter(estado=estado)
    if dpto_id:
        qs = qs.filter(distrito__provincia__departamento__id_ubigeo=dpto_id)

    # Ordenar por apellido paterno
    qs = qs.order_by('apellido_paterno')

    # Generar la respuesta HTTP del CSV/Excel
    import csv
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    filename = f"personeros_export_{timezone.now().strftime('%Y%m%d_%H%M%S')}.csv"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    # Escribir el BOM de UTF-8 para que Excel lo abra con formato correcto
    response.write(b'\xef\xbb\xbf')

    writer = csv.writer(response, delimiter=';') # Punto y coma es el estándar de Excel en español

    # Encabezados
    writer.writerow([
        'DNI',
        'Apellido Paterno',
        'Apellido Materno',
        'Nombres',
        'Celular',
        'Fecha Registro',
        'Departamento',
        'Provincia',
        'Distrito',
        'Centro de Votación',
        'Mesa',
        'Cargo',
        'Estado',
        'Perfil Completado',
        'Fecha Completado',
        'Observaciones'
    ])

    for p in qs:
        # Resolver relaciones de forma segura para evitar AttributeError si hay campos nulos
        dpto = p.distrito.provincia.departamento.nombre if p.distrito and p.distrito.provincia and p.distrito.provincia.departamento else 'Sin asignar'
        prov = p.distrito.provincia.nombre if p.distrito and p.distrito.provincia else 'Sin asignar'
        dist = p.distrito.nombre if p.distrito else 'Sin asignar'
        cv = p.centro_votacion.nombre if p.centro_votacion else 'Sin asignar'

        fecha_reg = timezone.localtime(p.fecha_creacion).strftime('%d/%m/%Y %H:%M:%S') if p.fecha_creacion else ''
        fecha_comp = timezone.localtime(p.fecha_completado).strftime('%d/%m/%Y %H:%M:%S') if p.fecha_completado else ''
        perfil_comp = 'Sí' if p.perfil_completado else 'No'

        writer.writerow([
            p.dni,
            p.apellido_paterno,
            p.apellido_materno,
            p.nombres,
            p.nro_celular,
            fecha_reg,
            dpto,
            prov,
            dist,
            cv,
            p.numero_mesa,
            p.get_cargo_display(),
            p.get_estado_display(),
            perfil_comp,
            fecha_comp,
            p.observaciones
        ])

    return response


@login_required(login_url='personeros:login')
def descargar_credencial_view(request, pk):
    """Genera y descarga la credencial del personero en formato PDF."""
    import os
    import platform
    import subprocess
    import docx
    from django.http import Http404

    # 1. Obtener el personero
    personero = get_object_or_404(Personero, pk=pk)

    # 2. Verificar permisos
    perfil = get_perfil(request.user)
    is_owner = False
    try:
        is_owner = (request.user.personero.pk == personero.pk)
    except Exception:
        pass

    # Permitir si es superusuario, si no tiene perfil (se asume admin en la app),
    # o si su perfil es de admin/coordinador, o si es el propio personero.
    has_permission = (
        request.user.is_superuser or
        not perfil or
        perfil.es_superadmin or
        perfil.es_coordinador or
        is_owner
    )

    if not has_permission:
        return HttpResponse("No tienes permisos para descargar esta credencial.", status=403)

    # 3. Ruta de la plantilla docx
    template_path = os.path.join(settings.BASE_DIR, 'docs', 'Credencial personero centro de votación.docx')
    if not os.path.exists(template_path):
        raise Http404("La plantilla de credencial no se encuentra en el servidor.")

    temp_docx_path = None
    temp_pdf_path = None

    try:
        # Helper para formatear, alinear y centrar vertical y horizontalmente las celdas
        def set_cell_text(cell, text, bold=False, size_pt=None):
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.shared import Pt

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if not cell.paragraphs:
                p = cell.add_paragraph()
            else:
                p = cell.paragraphs[0]

            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run(text)
            if bold:
                run.bold = True
            if size_pt:
                run.font.size = Pt(size_pt)
            run.font.name = 'Arial'

        # 4. Modificar el archivo docx con los datos del personero
        doc = docx.Document(template_path)

        # Configurar la hoja en formato A4 (21.0 cm x 29.7 cm) para compatibilidad estándar
        from docx.shared import Cm
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # Obtener referencias a las tablas originales del documento antes de crear la tabla superior
        table1 = doc.tables[0]
        table2 = doc.tables[1]

        # Llenar Tabla 1 (Nombres y DNI) - Centrado, Negrita y Letra más grande (12pt)
        set_cell_text(table1.rows[0].cells[1], personero.nombre_completo.upper(), bold=True, size_pt=12)
        set_cell_text(table1.rows[1].cells[1], personero.dni, bold=True, size_pt=12)

        # Generar código QR de seguridad dinámico en memoria
        import qrcode
        import io
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        if not personero.token_uuid:
            import uuid
            personero.token_uuid = uuid.uuid4()
            personero.save(update_fields=['token_uuid'])

        # Generar URL de validación pública
        qr_url = f"https://{request.get_host()}/personeros/validar/credencial/{personero.token_uuid}/"
        
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=1,
        )
        qr.add_data(qr_url)
        qr.make(fit=True)

        qr_img = qr.make_image(fill_color="black", back_color="white")
        qr_buffer = io.BytesIO()
        qr_img.save(qr_buffer, format="PNG")
        qr_buffer.seek(0)

        # Obtener referencias a los párrafos originales ANTES de modificar el documento
        p0_elem = doc.paragraphs[0]._element
        p1_elem = doc.paragraphs[1]._element

        # Crear una tabla invisible de maquetación de 1x3 en la parte superior para centrar los títulos y alinear el QR a la derecha
        table_top = doc.add_table(rows=1, cols=3)
        table_top.autofit = False

        # Mover la tabla al inicio absoluto de la estructura del documento
        tbl = table_top._tbl
        doc.element.body.remove(tbl)
        doc.element.body.insert(0, tbl)

        cell_left = table_top.rows[0].cells[0]
        cell_middle = table_top.rows[0].cells[1]
        cell_right = table_top.rows[0].cells[2]

        # Configurar anchos proporcionales alineados a la perfección con la Tabla 1 (total 6.15 pulgadas)
        # Al tener anchos iguales a la izquierda y derecha, la celda del centro queda perfectamente centrada
        cell_left.width = Inches(0.90)
        cell_middle.width = Inches(4.35)
        cell_right.width = Inches(0.90)

        # Configurar explícitamente los anchos en la cuadrícula de la tabla para asegurar que Word y LibreOffice los respeten
        for i, col in enumerate(table_top.columns):
            col.width = [Inches(0.90), Inches(4.35), Inches(0.90)][i]

        # Mover los párrafos de título originales a la celda del medio
        p0_elem.getparent().remove(p0_elem)
        p1_elem.getparent().remove(p1_elem)

        cell_middle._tc.append(p0_elem)
        cell_middle._tc.append(p1_elem)

        # Eliminar el párrafo vacío por defecto en la celda del medio
        default_p_mid = cell_middle.paragraphs[0]._element
        default_p_mid.getparent().remove(default_p_mid)

        # Limpiar el párrafo por defecto en la celda izquierda y hacerlo de tamaño mínimo para no agregar altura extra
        p_left = cell_left.paragraphs[0]
        p_left.text = ""
        p_left.paragraph_format.space_before = Pt(0)
        p_left.paragraph_format.space_after = Pt(0)
        p_left.paragraph_format.line_spacing = Pt(1)

        # Centrar ambos párrafos en la celda del medio, para que queden perfectamente alineados
        cell_middle.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_middle.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Insertar el QR en la celda derecha centrado verticalmente y alineado a la derecha
        cell_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if not cell_right.paragraphs:
            p_qr = cell_right.add_paragraph()
        else:
            p_qr = cell_right.paragraphs[0]
        p_qr.text = ""
        p_qr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_qr.paragraph_format.space_before = Pt(0)
        p_qr.paragraph_format.space_after = Pt(0)

        run_qr = p_qr.add_run()
        run_qr.add_picture(qr_buffer, width=Inches(0.9), height=Inches(0.9))

        # Llenar Tabla 2 (Local de Votación y ubicación) - Centrado utilizando la referencia original
        cv_nombre = personero.centro_votacion.nombre.upper() if personero.centro_votacion else "SIN ASIGNAR"
        cv_direccion = personero.centro_votacion.direccion.upper() if personero.centro_votacion and personero.centro_votacion.direccion else "SIN DIRECCIÓN"

        dist = personero.distrito.nombre.upper() if personero.distrito else "—"
        prov = personero.provincia.nombre.upper() if personero.distrito and personero.distrito.provincia else "—"
        dpto = personero.departamento.nombre.upper() if personero.departamento else "—"
        ubicacion = f"{dist}, {prov}, {dpto}"

        set_cell_text(table2.rows[0].cells[1], cv_nombre, bold=False, size_pt=11)
        set_cell_text(table2.rows[1].cells[1], cv_direccion, bold=False, size_pt=10)
        set_cell_text(table2.rows[2].cells[1], ubicacion, bold=False, size_pt=11)

        # Ajustar tamaños de letra: Firma PERSONERO LEGAL ANTE EL JEE DEL CALLAO (a 8pt)
        # y reglamento de personeros (a 7pt)
        en_reglamento = False
        for p in doc.paragraphs:
            text_upper = p.text.upper()
            if "PERSONERO LEGAL ANTE EL JEE DEL CALLAO" in text_upper:
                for run in p.runs:
                    run.font.size = Pt(8.0)

            if "REGLAMENTO SOBRE LA PARTICIPACI" in text_upper:
                en_reglamento = True

            if en_reglamento:
                for run in p.runs:
                    run.font.size = Pt(7.0)

        # 5. Crear nombres de archivos temporales
        temp_dir = os.path.join(settings.BASE_DIR, 'temp_credentials')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        safe_dni = personero.dni
        temp_docx_filename = f"credencial_{safe_dni}.docx"
        temp_pdf_filename = f"credencial_{safe_dni}.pdf"

        temp_docx_path = os.path.join(temp_dir, temp_docx_filename)
        temp_pdf_path = os.path.join(temp_dir, temp_pdf_filename)

        # Guardar docx modificado
        doc.save(temp_docx_path)

        # 6. Convertir a PDF con LibreOffice
        soffice_path = getattr(settings, 'SOFFICE_PATH', None)
        if not soffice_path:
            if platform.system() == 'Windows':
                soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
            else:
                soffice_path = "soffice"

        # Verificar existencia si es Windows para dar un mensaje claro
        if platform.system() == 'Windows' and not os.path.exists(soffice_path):
            return HttpResponse(
                "Error en el servidor: LibreOffice no se encuentra instalado en la ruta estándar de Windows.",
                status=500
            )

        command = [
            soffice_path,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            temp_dir,
            temp_docx_path
        ]

        # Ejecutar la conversión de forma síncrona
        result = subprocess.run(command, capture_output=True, text=True, timeout=30)

        if result.returncode != 0 or not os.path.exists(temp_pdf_path):
            return HttpResponse(
                f"Error al generar el PDF: {result.stderr or 'Error desconocido en LibreOffice'}",
                status=500
            )

        # 7. Leer el archivo PDF generado para retornarlo
        with open(temp_pdf_path, 'rb') as f:
            pdf_data = f.read()

        # 8. Responder con el archivo PDF para descarga
        response = HttpResponse(pdf_data, content_type='application/pdf')
        clean_name = personero.nombre_completo.replace(' ', '_').lower()
        response['Content-Disposition'] = f'attachment; filename="credencial_{clean_name}_{safe_dni}.pdf"'
        return response

    except Exception as e:
        return HttpResponse(f"Ocurrió un error inesperado al procesar la credencial: {str(e)}", status=500)
    finally:
        # 9. Limpieza de archivos temporales
        try:
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                os.remove(temp_pdf_path)
        except Exception:
            pass


def validar_credencial_publica_view(request, token_uuid):
    """Vista pública para validar una credencial mediante código QR sin requerir inicio de sesión."""
    personero = get_object_or_404(Personero, token_uuid=token_uuid)
    
    # La credencial se considera válida si el personero está confirmado y su perfil se completó
    es_valida = (personero.estado == 'confirmado')
    
    context = {
        'personero': personero,
        'es_valida': es_valida,
    }
    return render(request, 'personeros/validar_credencial.html', context)


def handler404_redirect(request, exception=None):
    """Redirige errores 404 al login, excepto para archivos estáticos."""
    if request.path.startswith(settings.STATIC_URL) or '/static/' in request.path:
        from django.http import HttpResponseNotFound
        return HttpResponseNotFound("Archivo estático no encontrado.")
    return redirect('personeros:login')
